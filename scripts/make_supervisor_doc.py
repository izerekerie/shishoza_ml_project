"""Builds TreeSight_Supervisor_Notes.docx — one clean document (opens in Google Docs)
combining the plain-language explanation of the whole pipeline + real code snippets.
Written in first person ('I') so Kerie can present it directly."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

GREEN = RGBColor(0x1B, 0x5E, 0x20)
GREY  = RGBColor(0x55, 0x55, 0x55)
CODEBG = "F2F2F2"

doc = Document()

# base style
st = doc.styles['Normal'].font
st.name = 'Calibri'; st.size = Pt(11)


def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = GREEN
    p.space_before = Pt(10); p.space_after = Pt(4)


def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = GREEN
    p.space_before = Pt(8); p.space_after = Pt(2)


def body(text, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic
    if color: r.font.color.rgb = color
    p.space_after = Pt(4)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def code(snippet, caption=None):
    if caption:
        c = doc.add_paragraph()
        r = c.add_run(caption); r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY
        c.space_after = Pt(1)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    for line in snippet.strip("\n").split("\n"):
        run = p.add_run(line + "\n")
        run.font.name = 'Consolas'; run.font.size = Pt(9)
    # shade the paragraph
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), CODEBG); pPr.append(shd)
    p.space_after = Pt(6)


# ───────────────────────────── TITLE
t = doc.add_paragraph()
r = t.add_run("TreeSight — Supervisor Meeting Notes")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = GREEN
sub = doc.add_paragraph()
r = sub.add_run("Detecting small-scale forest clearing in Rwanda's Nyungwe buffer zone  ·  "
                "Random Forest on Sentinel-2 + Sentinel-1 + SRTM, Hansen labels")
r.italic = True; r.font.color.rgb = GREY
doc.add_paragraph()

# ───────────────────────────── 1. DATA TYPES
h1("1. My three data types — what each does")

h2("Optical (Sentinel-2) — the colour camera")
bullet("Sees colour + near-infrared, which I use to compute NDVI (greenness).")
bullet("Weakness: blind through clouds — bad in Rwanda's two rainy seasons.")
bullet("On its own (Experiment A): F1 = 0.698.")

h2("Terrain (SRTM) — the height map")
bullet("Elevation, from which I compute slope (steepness) and aspect (hill direction).")
bullet("Helps because Nyungwe is steep and fragmented; clearing follows the terrain.")
bullet("Optical + Terrain (Experiment B): F1 = 0.777 -> adds +0.079 (my biggest jump).")

h2("Radar (Sentinel-1) — the echo sensor")
bullet("Sends microwave pulses and measures the bounce; senses structure / roughness.")
bullet("Strength: sees through clouds, day or night -> fills the rainy-season gap.")
bullet("Optical + Radar (Experiment C): F1 = 0.736 -> adds +0.038.")

h2("All combined (Experiment D)")
bullet("17 features together: F1 = 0.791 -> +0.093 over optical alone.")
bullet("Beats the published baseline of 0.71.")

h2("How exactly each one helps (summary)")
bullet("Terrain helped most (+0.079): in steep Nyungwe, slope/elevation carry strong signal.")
bullet("Radar helped less (+0.038) but matters for cloud cover, not raw score.")
bullet("Combining all three is best: each covers the others' weakness.")

# ───────────────────────────── 2. DATA COLLECTION
h1("2. How I get the data")
bullet("I do not take photos or screenshots.")
bullet("Satellites already capture imagery continuously; Google Earth Engine stores it.")
bullet("My script: pick Nyungwe + dates -> remove clouds -> build features -> export a table.")

# ───────────────────────────── 3. HANSEN
h1("3. My labels — Hansen Global Forest Change")
bullet("Hansen = my answer key (each pixel: cleared vs not cleared).")
bullet("Chose it: peer-reviewed (Science), free, yearly, reproducible.")
bullet("Same source as the baseline -> fair comparison.")

h2("Anticipated question: Hansen is 30 m but my features are 10 m — what's the point?")
bullet("There is no free, multi-year, Rwanda-wide 10 m deforestation label; Hansen is the only "
       "credible, reproducible option (hand-labelling thousands of pixels in 10 weeks is impossible).")
bullet("The label says WHERE loss happened; my 10 m features still add finer detail "
       "(NDVI-change, radar texture). Earth Engine resamples the 30 m label onto the 10 m grid.")
bullet("Honest limitation: 0.18 ha is about 2 Hansen pixels, so at the smallest scale the label "
       "itself is imprecise — some 'errors' may be label errors. This coarseness is part of why "
       "Rwanda's small clearing is underreported (ties to my motivation).")
body("Alternatives I considered and why I did not use them:")
bullet("Hand-labelling from high-res (Google/Planet): too slow, subjective, not reproducible.")
bullet("RADD alerts (~10 m): built for humid-tropics alerts / primary forest, not Rwanda smallholder labels.")
bullet("Dynamic World (10 m): land-cover, not a forest-loss change product; noisy year-to-year.")
bullet("Field GPS ground truth: impossible at scale (FMES has one tablet per district).")
body("Future work: refine labels with RADD / Dynamic World, and manually verify a sample of "
     "Hansen labels against high-resolution imagery.", color=GREY)

# ───────────────────────────── 4. MATCHING
h1("4. How my features match the label")
bullet("1 row = 1 pixel = one 10 m spot on the ground.")
bullet("17 feature columns = the measurements at that spot.")
bullet("Label column = Hansen's answer at the SAME spot (1 = cleared 2020-22, 0 = forest).")
bullet("They match by location — Earth Engine stacks all layers on one grid and samples each point "
       "(like poking a pin through a stack of transparent map sheets).")

# ───────────────────────────── 5. MODEL
h1("5. My model — Random Forest")
bullet("Many decision trees; each asks yes/no questions on the features -> guesses cleared / not.")
bullet("All trees vote; majority wins (accurate and stable).")
body("Why I chose it:")
bullet("Works on small data (deep learning needs far more — my lit review confirms).")
bullet("Runs on a normal laptop, no special hardware.")
bullet("Gives feature importance -> directly answers RQ1 (does radar help?).")

# ───────────────────────────── 6. COMPARISON
h1("6. How I compare with existing models")
bullet("Same metric (F1) for fairness.")
bullet("Baseline = 0.71; mine = 0.79 random split (and the stricter spatial number — see section 9).")
bullet("Other studies (Ghana, E. Rwanda): different areas -> compare methods, not raw numbers.")
bullet("None detect at 0.18 ha or give citizens an interface -> that gap is mine.")

# table
tbl = doc.add_table(rows=1, cols=5); tbl.style = 'Light Grid Accent 1'
hdr = tbl.rows[0].cells
for i, txt in enumerate(["Model / Tool", "Data", "Smallest reliable", "F1 / acc", "Citizen UI?"]):
    hdr[i].paragraphs[0].add_run(txt).bold = True
for row in [
    ["GFW / RADD (Ygorra 2024)", "Optical/radar", "0.1-6.25 ha", "F1 ~0.71", "No"],
    ["Amoakoh 2021 (Ghana)", "Opt+radar+terrain RF", "-", "94% acc", "No"],
    ["Gutkin 2023 (E. Rwanda)", "Optical-only RF", "-", "85-92%", "No"],
    ["Mine (TreeSight, Nyungwe)", "Opt+radar+terrain RF", "~0.1 ha (rec 0.82)", "F1 0.79", "Yes"],
]:
    cells = tbl.add_row().cells
    for i, v in enumerate(row):
        cells[i].paragraphs[0].add_run(v)

# ───────────────────────────── 7. RESULTS RQ1
h1("7. My results — RQ1 (does radar / terrain help?)")
bullet("Optical only: F1 = 0.698.")
bullet("Optical + Terrain: F1 = 0.777 (+0.079).")
bullet("Optical + Radar: F1 = 0.736 (+0.038).")
bullet("All combined: F1 = 0.791 -> beats the baseline of 0.71.")
bullet("Finding: terrain helped most; combining all three is best.")

# ───────────────────────────── 8. RESULTS RQ2
h1("8. My results — RQ2 (how small can it detect?)")
bullet("<= 0.1 ha (tiny): recall = 0.82.")
bullet("0.1-0.2 ha (the 0.18 ha smallholder case): recall = 0.80.")
bullet("0.2-0.5 ha: recall = 0.83.")
bullet("Finding: recall stays ~0.80 even for the tiniest patches — it does not collapse like global tools.")
bullet("Honest caveat: very few large patches to compare against, so this comparison is limited.")

# ───────────────────────────── 9. TESTING
h1("9. How I tested (my honesty layer)")
bullet("Balanced data, stratified 80/20 split.")
bullet("Checked for label leakage -> found 0.")
bullet("Checked how far test pixels sit from training pixels (median 218 m).")
bullet("Random split: F1 = 0.79. Other metrics: precision 0.77, recall 0.82, AUC 0.87.")
bullet("Confusion matrix: 758 / 817 correct, 242 / 183 wrong.")

h2("What spatial cross-validation means")
bullet("Problem: nearby pixels look almost the same, so a random split can be too optimistic.")
bullet("Random split mixes train/test pixels on the map (neighbours) -> inflated F1 = 0.79.")
bullet("Spatial CV splits by LOCATION (blocks/regions): test on ground the model never saw "
       "-> the honest, real-world number.")

h2("My spatial CV results (computed and saved)")
bullet("Spatial F1 depends on block size — I report the range, not one number:")
bullet("5 km blocks: 0.768  |  8 km: 0.765  |  11 km: 0.754  |  15 km: 0.743  |  10 KMeans blocks: 0.733.")
bullet("Random split, for reference: 0.79.")
bullet("The pattern IS a finding: bigger blocks (more train/test separation) -> lower F1 = "
       "spatial autocorrelation, measured on my own data.")
bullet("Honest headline: spatial F1 = 0.73-0.77 (about 0.75 at 11 km blocks). EVERY variant still "
       "beats the 0.71 baseline — and the spatial test is stricter than the baseline's.")
body("Saved to results/metrics/spatial_cv_results.json.", color=GREY)

h2("How F1 measures accuracy")
bullet("Precision = when I say 'cleared', how often am I right? 817/1059 = 0.77 (fewer false alarms).")
bullet("Recall = of all real clearings, how many did I catch? 817/1000 = 0.82 (fewer misses).")
bullet("F1 = balance of both = 2*(P*R)/(P+R) = 0.79.")
bullet("Why not plain accuracy: if most land is forest, a lazy 'never cleared' model scores high "
       "while catching nothing. F1 focuses on the rare clearing class — exactly my goal.")

# ───────────────────────────── 10. TUNING
h1("10. Hyperparameter tuning")
bullet("Grid search: 96 combinations x 5-fold CV = 480 model fits.")
bullet("Best settings: 800 trees, max depth 25, max_features sqrt, min_samples_leaf 1.")
bullet("Finding: tuning barely changed F1 (0.794 -> 0.791) -> my data and features matter more "
       "than the settings.")

# ───────────────────────────── 11. INTERESTING
h1("11. What makes my research interesting")
bullet("Beat the global baseline by training locally (0.71 -> 0.79).")
bullet("Terrain mattered more than radar in steep Nyungwe — surprising.")
bullet("Small patches not badly missed (~0.82 recall).")
bullet("Tuning barely helped -> features dominate.")
bullet("Honest gap between random (0.79) and spatial test.")

# ───────────────────────────── 12. NATIONAL EXTENSION
h1("12. National extension — I tested training on all 5 provinces")
body("I kept my Nyungwe pipeline intact and added a national dataset (23,319 pixels across all "
     "five provinces) as an extension, to test whether training nationally helps. It does — on "
     "every measure.")

h2("National beats Nyungwe-only on all four experiments (5-fold CV)")
t1 = doc.add_table(rows=1, cols=4); t1.style = 'Light Grid Accent 1'
for i,txt in enumerate(["Experiment","Nyungwe F1","National F1","Change"]):
    t1.rows[0].cells[i].paragraphs[0].add_run(txt).bold = True
for r in [["A — Optical only","0.697","0.755","+0.058"],
          ["B — Optical + Terrain","0.766","0.821","+0.054"],
          ["C — Optical + Radar","0.729","0.794","+0.064"],
          ["D — All combined","0.783","0.832","+0.049"]]:
    cells=t1.add_row().cells
    for i,v in enumerate(r): cells[i].paragraphs[0].add_run(v)

h2("It generalizes across regions (leave-one-province-out)")
bullet("Held out a WHOLE province the model never saw, then tested on it:")
bullet("Kigali 0.850 | North 0.806 | South 0.806 | East 0.805 | West 0.709.")
bullet("Even unseen, F1 stays 0.71-0.85 -> this is direct evidence for 'citizens anywhere'.")
bullet("West is hardest (0.709) — the steep montane zone Nyungwe sits in, which explains why "
       "Nyungwe-only was a tough baseline.")
bullet("National spatial-CV F1 = 0.753 (vs Nyungwe 0.733).")

h2("RQ findings on national data")
bullet("RQ1: by source optical 0.51, radar 0.26, terrain 0.23 — nationally radar slightly "
       "out-ranks terrain (opposite of Nyungwe, where terrain led). That contrast is a finding.")
bullet("RQ2: small-patch detection improves — the 0.18 ha case rises from 0.80 to 0.88 recall.")

h2("Honest caveats")
bullet("National data is bigger (23k vs 10k), so some gain is more data, not only diversity.")
bullet("Labels are still Hansen 30 m; spatial variance is higher (region-dependent, +/-0.08).")

# ───────────────────────────── 13. GUIDANCE
h1("13. Where I want your guidance")
bullet("Scope (my main question): my national results now IMPROVE accuracy (+~0.05) AND "
       "generalize across all provinces (0.71-0.85), and the national data already exists. So I "
       "now lean toward going NATIONAL with Nyungwe as my primary case study — does this hold, or "
       "do you prefer I keep the scope Nyungwe-only?")
bullet("Is my spatial-CV F1 (0.73-0.77 range) a defensible headline alongside the random 0.79?")
bullet("For RQ2, is recall-by-patch-size enough, or do you also want precision by size?")
bullet("Is 10 parcels a defensible app-validation sample?")

# ───────────────────────────── APPENDIX: CODE
doc.add_page_break()
h1("Appendix — Key code (the real snippets)")

h2("A. Hansen labels — feature extraction (Google Earth Engine)")
code("""
var hansen = ee.Image('UMD/hansen/global_forest_change_2023_v1_11').clip(nyungwe);

// Pixels that lost forest 2020-2022 = training label 1 (lossyear 20..22)
var loss_train = hansen.select('lossyear').gte(20)
  .and(hansen.select('lossyear').lte(22)).rename('label');

// Stable forest = treecover2000 >= 30% AND no loss = label 0
// .unmask(0) is critical so stable pixels are sampled, not masked out
var stable_forest = hansen.select('treecover2000').gte(30)
  .and(hansen.select('loss').eq(0)).unmask(0);

// All feature layers stacked on ONE grid, then the integer label restored
var stack = ee.Image.cat([
  ndvi_train, ndvi_test, ndvi_change,       // optical
  s1_train.select('VH'), s1_train.select('VV'), radar_ratio,  // radar
  elevation, slope, aspect                  // terrain
]).float().addBands(label_train, ['label'], true);
""", caption="notebooks/01_GEE_Export.js — each sampled pixel pulls all features + the Hansen label from the same location")

h2("B. Random Forest — the four experiments")
code("""
model = RandomForestClassifier(
    n_estimators=200,          # 200 trees vote together
    max_depth=20,
    min_samples_leaf=5,
    class_weight='balanced',   # handles forest/cleared imbalance
    random_state=42,           # reproducible
    n_jobs=-1                   # use all CPU cores
)
model.fit(X_train, y_train)
f1 = f1_score(y_test, model.predict(X_test))
""", caption="notebooks/03_Train_Model.ipynb — run once per feature set A/B/C/D")

h2("C. Grid search — and why I trust the chosen parameters")
code("""
param_grid = {
    "n_estimators":     [100, 200, 400, 800],
    "max_depth":        [None, 10, 15, 25],
    "min_samples_leaf": [1, 2, 5],
    "max_features":     ["sqrt", "log2"],
}   # 4 x 4 x 3 x 2 = 96 combos

gs = GridSearchCV(rf, param_grid, scoring='f1', cv=5, n_jobs=-1)
gs.fit(X_train_D, y_train)     # 96 x 5 = 480 fits
print(gs.best_params_)         # -> 800 trees, depth 25, sqrt, leaf 1
""", caption="scripts/hyperparameter_tune.py — parameters are CHOSEN by 5-fold CV, not guessed. Result: tuning barely beat the defaults, so features matter more than settings.")

h2("D. Spatial leakage check — what I have now")
code("""
from scipy.spatial import cKDTree
# distance from each TEST pixel to its nearest TRAINING pixel
tree = cKDTree(train_coords)            # lat/lng of train pixels
nearest_deg, _ = tree.query(test_coords, k=1)
# result: median ~218 m, but 18% within 60 m  ->  possible optimism
""", caption="scripts/evaluate_split_and_patchsize.py — this CHECKS leakage; it does not yet RUN spatial CV")

h2("E. Spatial cross-validation — the code I ran (result: F1 0.73-0.77)")
code("""
import numpy as np, pandas as pd, json
from sklearn.cluster import KMeans
from sklearn.model_selection import GroupKFold
from sklearn.metrics import f1_score

# coordinates come from the raw GEE export (.geo), row-aligned with the clean data
lat = raw['.geo'].apply(lambda s: json.loads(s)['coordinates'][1]).values
lng = raw['.geo'].apply(lambda s: json.loads(s)['coordinates'][0]).values

# turn coordinates into spatial blocks, then hold out WHOLE blocks
blocks = KMeans(n_clusters=10, random_state=42).fit_predict(np.column_stack([lat, lng]))
gkf, f1s = GroupKFold(n_splits=5), []
for tr, te in gkf.split(X, y, groups=blocks):
    m = RandomForestClassifier(n_estimators=800, max_depth=25, min_samples_leaf=1,
            max_features='sqrt', class_weight='balanced',
            random_state=42, n_jobs=-1).fit(X[tr], y[tr])
    f1s.append(f1_score(y[te], m.predict(X[te])))

print('Spatial-CV F1:', round(np.mean(f1s), 3))   # 0.733 (10 KMeans blocks)
# Repeating with 5/8/11/15 km grid cells gives 0.768 / 0.765 / 0.754 / 0.743
""", caption="My computed spatial result, saved to results/metrics/spatial_cv_results.json. Bigger blocks -> lower F1, all above the 0.71 baseline.")

out = "TreeSight_Supervisor_Notes.docx"
doc.save(out)
print("Saved", out)
